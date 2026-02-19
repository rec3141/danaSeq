#!/usr/bin/env python3
"""Convert the AI MAG Bin Refinement Plan markdown to a professional Word document."""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

INPUT_MD = "/home/grid/.claude/plans/polished-scribbling-lampson.md"
OUTPUT_DOCX = "/data/danav2/nanopore_mag/AI_MAG_Bin_Refinery_Plan.docx"

HEADING_COLOR = RGBColor(0x1A, 0x3C, 0x5E)
BODY_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
SUBTITLE_COLOR = RGBColor(0x4A, 0x4A, 0x4A)
CODE_BG = "F0F0F0"
CODE_TEXT_COLOR = RGBColor(0x1A, 0x1A, 0x2E)
INLINE_CODE_COLOR = RGBColor(0x80, 0x00, 0x00)
TABLE_HEADER_BG = "2E4057"
TABLE_ALT_ROW_BG = "F5F5F5"


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_run(paragraph, text, bold=False, italic=False, code=False):
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = INLINE_CODE_COLOR
    return run


def parse_inline_formatting(paragraph, text):
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*'
        r'|\*\*(.+?)\*\*'
        r'|\*(.+?)\*'
        r'|`([^`]+?)`)'
    )
    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            add_formatted_run(paragraph, text[last_end:match.start()])
        if match.group(2):
            add_formatted_run(paragraph, match.group(2), bold=True, italic=True)
        elif match.group(3):
            add_formatted_run(paragraph, match.group(3), bold=True)
        elif match.group(4):
            add_formatted_run(paragraph, match.group(4), italic=True)
        elif match.group(5):
            add_formatted_run(paragraph, match.group(5), code=True)
        last_end = match.end()
    if last_end < len(text):
        add_formatted_run(paragraph, text[last_end:])


def add_code_block(doc, code_lines):
    for line in code_lines:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = CODE_TEXT_COLOR
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
        pPr.append(shading)


def add_table_from_rows(doc, header, rows):
    if not header or not rows:
        return
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, TABLE_HEADER_BG)
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(min(n_cols, len(row_data))):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            parse_inline_formatting(p, row_data[c_idx].strip())
            for run in p.runs:
                if not run.font.name:
                    run.font.name = "Calibri"
                if not run.font.size:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_ROW_BG)
    doc.add_paragraph()


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_COLOR
    sizes = {1: 24, 2: 18, 3: 14, 4: 12}
    for level in range(1, 5):
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.name = 'Calibri'
        hstyle.font.color.rgb = HEADING_COLOR
        hstyle.font.size = Pt(sizes[level])

    # Title block
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    title_para.paragraph_format.space_after = Pt(4)
    run = title_para.add_run("AI MAG Bin Refinement Swarm")
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = "Calibri"
    run.font.color.rgb = HEADING_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run("Architecture Design")
    run.font.size = Pt(18)
    run.font.name = "Calibri"
    run.font.color.rgb = SUBTITLE_COLOR

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(24)
    pPr = rule._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="1A3C5E"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    with open(md_path, 'r') as f:
        lines = f.readlines()

    i = 0
    if lines and lines[0].startswith('# '):
        i = 1

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Fenced code blocks
        if line.strip().startswith('```'):
            code_lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            if code_lang:
                lbl = doc.add_paragraph()
                lbl.paragraph_format.space_after = Pt(0)
                lbl.paragraph_format.space_before = Pt(8)
                r = lbl.add_run(code_lang.upper())
                r.font.size = Pt(8)
                r.font.name = "Calibri"
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                r.bold = True
            add_code_block(doc, code_lines)
            if i < len(lines):
                i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            h = doc.add_heading(level=level)
            parse_inline_formatting(h, heading_text)
            for run in h.runs:
                run.font.name = "Calibri"
                run.font.color.rgb = HEADING_COLOR
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                header_cells = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                data_start = 1
                if re.match(r'^[\|\s\-:]+$', table_lines[1]):
                    data_start = 2
                data_rows = []
                for tl in table_lines[data_start:]:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    data_rows.append(cells)
                add_table_from_rows(doc, header_cells, data_rows)
            continue

        # Blockquotes
        if line.strip().startswith('>'):
            bq_text = line.strip().lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:left w:val="single" w:sz="18" w:space="4" w:color="4A90D9"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            parse_inline_formatting(p, bq_text)
            for run in p.runs:
                if not run.font.color.rgb:
                    run.font.color.rgb = SUBTITLE_COLOR
                run.font.italic = True
            i += 1
            continue

        # Bullet lists
        bullet_match = re.match(r'^(\s*)([-*+])\s+(.*)', line)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            bullet_text = bullet_match.group(3)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.27 + indent_level * 0.63)
            parse_inline_formatting(p, bullet_text)
            for run in p.runs:
                if not run.font.name:
                    run.font.name = "Calibri"
                if not run.font.size:
                    run.font.size = Pt(11)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if num_match:
            indent_level = len(num_match.group(1)) // 2
            list_text = num_match.group(3)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(1.27 + indent_level * 0.63)
            parse_inline_formatting(p, list_text)
            for run in p.runs:
                if not run.font.name:
                    run.font.name = "Calibri"
                if not run.font.size:
                    run.font.size = Pt(11)
            i += 1
            continue

        # Empty lines
        if line.strip() == '':
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        parse_inline_formatting(p, line)
        for run in p.runs:
            if not run.font.name:
                run.font.name = "Calibri"
            if not run.font.size:
                run.font.size = Pt(11)
        i += 1

    doc.save(docx_path)
    print(f"SUCCESS: Document saved to {docx_path}")


if __name__ == '__main__':
    convert_md_to_docx(INPUT_MD, OUTPUT_DOCX)
